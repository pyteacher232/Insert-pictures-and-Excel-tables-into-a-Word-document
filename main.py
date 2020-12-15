# docx module is used for reading and writing a word document.
import docx
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor
from tqdm import *
import openpyxl

# listdir of os is used for getting all file names under a directory.
from os import listdir
# isfile is used for checking if it is a file or directory and join is used for joining path segments.
from os.path import isfile, join
import numpy as np

# Path of your word file.
word_fname = "Word Report- Python Automation.docx"
# Path of your pictures.
pic_dir = "pics"

# Get all pictures in the pics directory.
all_pics = [f for f in listdir(pic_dir) if isfile(join(pic_dir, f))]

doc = docx.Document(word_fname)

# Get the width of your word document.
width = doc.sections[0].page_width
# Get the left margin of your word document.
left_margin = doc.sections[0].left_margin
# Get the right margin of your word document.
right_margin = doc.sections[0].right_margin

# Calculate the width of images to be inserted pairwise.
img_width = (width - left_margin - right_margin) / 2 * 0.95

# Get all paragraphs.
all_paras = doc.paragraphs

# Read excel file
input_dt = {}
input_fname = 'Completed Data Update - NatGas Paper.xlsx'

wb_obj = openpyxl.load_workbook(input_fname)
sheet_names = wb_obj.sheetnames

for i, sheet_name in enumerate(sheet_names):
    sheet = wb_obj[sheet_name]

    if sheet_name not in input_dt:
        input_dt[sheet_name] = []

    for row_index in range(0, sheet.max_row):
        row = [(sheet.cell(row_index+1, col_index+1).value, sheet.cell(row_index+1, col_index+1).font) for col_index in range(sheet.max_column)]
        input_dt[sheet_name].append(row)

# Move table after a paragraph
def move_table_after(table, paragraph):
    tbl, p = table._tbl, paragraph._p
    p.addnext(tbl)

# Draw a cell border of table
def set_cell_border(cell, **kwargs):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()

    # check for tag existnace, if none found, then create one
    tcBorders = tcPr.first_child_found_in("w:tcBorders")
    if tcBorders is None:
        tcBorders = OxmlElement('w:tcBorders')
        tcPr.append(tcBorders)

    # list over all available tags
    for edge in ('start', 'top', 'end', 'bottom', 'insideH', 'insideV'):
        edge_data = kwargs.get(edge)
        if edge_data:
            tag = 'w:{}'.format(edge)

            # check for tag existnace, if none found, then create one
            element = tcBorders.find(qn(tag))
            if element is None:
                element = OxmlElement(tag)
                tcBorders.append(element)

            # looks like order of attributes is important
            for key in ["sz", "val", "color", "space", "shadow"]:
                if key in edge_data:
                    element.set(qn('w:{}'.format(key)), str(edge_data[key]))


for para in all_paras:
    print(para.text)

    # Insert 'Backtesting_Plot' pictures
    if "Backtesting Plot" in para.text:
        # Delete "Backtesting_Plot".
        para.text = ''
        r = para.add_run()
        for p in all_pics:
            if "Backtesting Plot" in p:
                r.add_picture(join(pic_dir, p), width=img_width)
                all_pics.remove(p)

    # Insert 'Backtesting_Plot' pictures
    if "Benchmarking Plot" in para.text:
        # Delete "Backtesting_Plot".
        para.text = ''
        r = para.add_run()
        for p in all_pics:
            if "Benchmarking Plot" in p:
                r.add_picture(join(pic_dir, p), width=img_width)
                all_pics.remove(p)

    # Insert 'Sensitivity_Plot' pictures
    if "Sensitivity Plot" in para.text:
        # Delete "Sensitivity_Plot".
        para.text = ''
        r = para.add_run()
        for p in all_pics:
            if "Sensitivity Plot" in p:
                r.add_picture(join(pic_dir, p), width=img_width)
                all_pics.remove(p)

    # Insert from Picture1 to Picture6.
    for i in range(6):
        if "Insert Picture{} here".format(i + 1) in para.text:
            # Delete "Insert Picture here...".
            para.text = ''

            # Add a picture.
            r = para.add_run()
            p = 'Picture{}.png'.format(i + 1)
            r.add_picture(join(pic_dir, p))
            all_pics.remove(p)

    # Insert the remaining pictures except from picture1 to picture6.
    if "Insert 16 pictures pairwise" in para.text:
        # Delete "Insert 16 pictures pairwise".
        para.text = ''
        r = para.add_run()

        # Add all the remaining pictures.
        for p in all_pics:
            r.add_picture(join(pic_dir, p), width=img_width)
            all_pics.remove(p)

    for sheet_name in input_dt:
        if sheet_name in para.text:
            para.text = ''
            # r = para.add_run()

            tbl_dt = input_dt[sheet_name]
            tbl_dt = np.array(tbl_dt)

            # tbl = doc.add_table(rows=tbl_dt.shape[0], cols=tbl_dt.shape[1], style=None)
            tbl = doc.add_table(rows=0, cols=tbl_dt.shape[1])
            # tbl.style = "Table Grid"
            move_table_after(tbl, para)

            bar = tqdm(total=tbl_dt.shape[0])
            bar.set_description(f"Inserting a table '{sheet_name}' now...")

            for row_idx in range(tbl_dt.shape[0]):
                cells = tbl.add_row().cells
                for col_idx in range(tbl_dt.shape[1]):
                    set_cell_border(
                        cells[col_idx],
                        top={"sz": 12, "val": "single", "color": "#000000", "space": "0"},
                        bottom={"sz": 12, "color": "#000000", "val": "single"},
                        start={"sz": 12, "val": "single", "color": "#000000", "shadow": "true"},
                        end={"sz": 12, "val": "single", "color": "#000000"},
                    )
                    cells[col_idx].text = str(tbl_dt[row_idx, col_idx, 0]) if tbl_dt[row_idx, col_idx, 0] else ""
                    for paragraph in cells[col_idx].paragraphs:
                        for run in paragraph.runs:
                            run.font.name = tbl_dt[row_idx, col_idx, 1].name
                            run.font.bold = tbl_dt[row_idx, col_idx, 1].b
                            run.font.size = Pt(int(tbl_dt[row_idx, col_idx, 1].sz))
                            run.font.italic = tbl_dt[row_idx, col_idx, 1].i


                            def hex_to_rgb(value):
                                value = value.lstrip('#')
                                lv = len(value)
                                return tuple(int(value[i:i + lv // 3], 16) for i in range(0, lv, lv // 3))

                            try:
                                run.font.color.rgb = RGBColor(*hex_to_rgb(tbl_dt[row_idx, col_idx, 1].color.rgb)[1:])
                            except:
                                pass

                bar.update()

            bar.close()

# Save the document
doc.save('Word Report- Python Automation_re.docx')
